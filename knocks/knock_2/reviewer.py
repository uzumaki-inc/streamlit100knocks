import json
from langchain.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from langchain_unstructured import UnstructuredLoader

from docx import Document
from docx.shared import RGBColor

# OpenAIの環境変数を読み込み
from common.util import load_environment
load_environment()

# WordファイルをHTMLに変換する関数
def word_to_html(file_path):
    doc = Document(file_path)
    html_content = ""

    for paragraph in doc.paragraphs:
        paragraph_html = ""

        for run in paragraph.runs:
            text = run.text
            style = ""

            # 太字チェック
            if run.bold:
                style += "font-weight:bold;"

            # 文字色チェック
            if run.font.color and run.font.color.rgb:
                color = run.font.color.rgb
                style += f"color:#{color};"

            # スタイルを適用してHTMLに変換
            paragraph_html += f"<span style='{style}'>{text}</span>"

        # 段落ごとに改行を追加
        html_content += f"<p>{paragraph_html}</p>"

    return html_content

# Wordファイルの内容を読み込む関数
def read_word_file(file_path):
    doc = Document(file_path)
    content = [paragraph.text for paragraph in doc.paragraphs]
    return "\n".join(content)

def load_word_file_with_langchain(file_path: str) -> str:
    """
    LangChainのUnstructuredLoaderを使用してWordファイルからテキストを抽出する。
    """
    loader = UnstructuredLoader(file_path)
    documents = loader.load()
    return "\n".join([doc.page_content for doc in documents])

def correct_text_with_llm(text: str) -> list:
    """
    テキストから誤字脱字や不適切な表現を修正し、JSON形式で返す関数。
    """
    llm = ChatOpenAI(model="gpt-4o", temperature=0)

    prompt = PromptTemplate(
        input_variables=["text"],
        template="""
        以下の日本語テキストに含まれる誤字脱字や不適切な表現を修正し、**必ず純粋なJSON形式**で出力してください。
        各修正箇所には、元の表現、修正後の表現、修正理由、行番号を含めてください
        JSONフォーマット:
        [
          {{
            "original": "<元の表現>",
            "corrected": "<修正後の表現>",
            "reason": "<修正理由>",
            "line_number": <行番号>
          }}
        ]

        出力するJSONは以下の条件を守ってください:
        1. **純粋なJSONオブジェクトのみ**を出力してください。
        2. コードブロック記号やコメントは一切含めないでください。
        3. テキスト以外の余分な空白や説明も含めないでください。

        テキスト:
        {text}

        注意: JSON形式以外のテキスト、余計な説明、追加のコメントは出力しないでください。出力は必ず**JSONオブジェクトのみ**にしてください。
        """
    )
    # RunnableLambdaでフォーマット済みのプロンプトを作成
    formatted_prompt = prompt.format(text=text)

    # ChatOpenAIに直接渡せる形式に変更
    response = llm.invoke(formatted_prompt)
    # AIMessageからcontentを取り出してJSONパース
    response_content = response.content  # ここでcontentを抽出
    print(response_content)

    try:
        corrections = json.loads(response_content)
    except json.JSONDecodeError:
        raise ValueError("LLMの出力がJSON形式ではありませんでした。")

    return corrections

def add_corrections_to_word(input_file: str, corrections: list, output_file: str):
    """
    Wordファイルに修正箇所を擬似コメントとして追加し、修正版を保存する。
    修正内容は対象段落の末尾に太字・赤色テキストとして追加される。
    """
    doc = Document(input_file)
    paragraphs = doc.paragraphs


    for correction in corrections:
        line_number = correction["line_number"] - 1  # 行番号を0ベースに変換
        original = correction["original"]
        corrected = correction["corrected"]
        reason = correction["reason"]

        # 指定された行番号の段落に修正案を追加
        if 0 <= line_number < len(paragraphs):
            paragraph = paragraphs[line_number]
            if original in paragraph.text:  # originalが含まれている場合のみ
                # 修正案を末尾に追加
                comment_text = f" [修正案: '{corrected}' 理由: {reason}]"
                run = paragraph.add_run(comment_text)
                run.bold = True  # 太字
                run.font.color.rgb = RGBColor(255, 0, 0)  # 赤色

    # Wordファイルを保存
    doc.save(output_file)
    print(f"修正版Wordファイルが {output_file} に保存されました。")

def process_word_file(input_file: str, output_file: str):
    """
    Wordファイルを読み込み、修正をコメントとして追加し、新しいWordファイルを保存する。
    """
    print("LangChainでWordファイルを読み込んでいます...")
    text = load_word_file_with_langchain(input_file)

    print("誤字脱字の修正を実行しています...")
    corrections = correct_text_with_llm(text)

    print("Wordファイルにコメントを追加しています...")
    add_corrections_to_word(input_file, corrections, output_file)

# 実行例
if __name__ == "__main__":
    input_word_file = "knocks/knock_2/input.docx"  # 入力Wordファイル
    output_word_file = "knocks/knock_2/output.docx"  # 出力Wordファイル

    process_word_file(input_word_file, output_word_file)
