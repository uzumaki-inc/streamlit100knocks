import os
from langchain.text_splitter import CharacterTextSplitter
from langchain_chroma import Chroma
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from docx import Document
from docx.shared import RGBColor
from common.util import load_environment
from langchain_core.output_parsers import JsonOutputParser
from langchain.schema.runnable import RunnablePassthrough
from operator import itemgetter


load_environment()

def create_review_chain():
    template = """
    あなたは日立のスタイルガイドに基づいて文章を校正する専門家です。

    スタイルガイドの関連部分:
    {context}

    以下の文章を校正してください:
    {text}

    必ず以下の形式のJSONで出力してください:
    [
        {{
            "original": "修正が必要な元の表現",
            "corrected": "修正後の表現",
            "reason": "修正理由",
            "line_number": 行番号
        }}
    ]
    """

    prompt = ChatPromptTemplate.from_template(template)
    return prompt

def review_text(text: str, db) -> list:
    model = ChatOpenAI(model="gpt-4o", temperature=0)
    prompt = create_review_chain()
    retriever = db.as_retriever()

    # レビューチェーンの構築
    review_chain = (
        {
            "text": RunnablePassthrough(),
            "context": retriever,
        }
        | prompt
        | model
        | JsonOutputParser()
    )

    # チェーンの実行
    try:
        result = review_chain.invoke(text)
        return result
    except Exception as e:
        print(f"エラー: {e}")
        return []

def add_corrections_to_word(input_file: str, corrections: list, output_file: str):
    """
    Wordファイルに修正箇所を擬似コメントとして追加し、修正版を保存する。
    修正内容は対象段落の末尾に太字・赤色テキストとして追加される。
    """
    doc = Document(input_file)
    paragraphs = doc.paragraphs


    for correction in corrections:
        line_number = correction["line_number"] - 1  # 行番号を0ベースに変換
        original = correction["original"]
        corrected = correction["corrected"]
        reason = correction["reason"]

        # 指定された行番号の段落に修正案を追加
        if 0 <= line_number < len(paragraphs):
            paragraph = paragraphs[line_number]
            if original in paragraph.text:  # originalが含まれている場合のみ
                # 修正案を末尾に追加
                comment_text = f" [修正案: '{corrected}' 理由: {reason}]"
                run = paragraph.add_run(comment_text)
                run.bold = True  # 太字
                run.font.color.rgb = RGBColor(255, 0, 0)  # 赤色

    # Wordファイルを保存
    doc.save(output_file)
    print(f"修正版Wordファイルが {output_file} に保存されました。")

def load_and_prepare_vectorstore(style_guide_path: str, persist_dir: str = "./chroma_db") -> Chroma:
    """
    スタイルガイドのベクトルストアを準備する。
    既存のベクトルストアがある場合は再利用し、ない場合は新規作成する。
    """
    # 既存のベクトルストアがあるか確認
    if os.path.exists(persist_dir):
        print("既存のベクトルストアを読み込んでいます...")
        embeddings = OpenAIEmbeddings()
        return Chroma(persist_directory=persist_dir, embedding_function=embeddings)

    # 新規作成の場合
    print("スタイルガイドを読み込んでいます...")
    with open(style_guide_path, "r", encoding="utf-8") as f:
        style_guide = f.read()

    print("新規ベクトルストアを作成しています...")
    text_splitter = CharacterTextSplitter(
        separator="\n\n",
        chunk_size=1000,
        chunk_overlap=200
    )
    chunks = text_splitter.split_text(style_guide)

    embeddings = OpenAIEmbeddings()

    # 本番環境ではメモリに保存する
    is_production = os.getenv("ENV") == "production"
    if is_production:
        persist_dir = None

    vectorstore = Chroma.from_texts(
        texts=chunks,
        embedding=embeddings,
        persist_directory=persist_dir
    )

    return vectorstore

def process_word_file(input_file: str, output_file: str):
    db = load_and_prepare_vectorstore(style_guide_path="style_guide.txt")

    print("Wordファイルを読み込んでいます...")
    doc = Document(input_file)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    print("文章を校正しています...")
    corrections = review_text(text, db)

    print("修正をWordファイルに適用しています...")
    add_corrections_to_word(input_file, corrections, output_file)

# WordファイルをHTMLに変換する関数
def word_to_html(file_path):
    doc = Document(file_path)
    html_content = ""

    for paragraph in doc.paragraphs:
        paragraph_html = ""

        for run in paragraph.runs:
            text = run.text
            style = ""

            # 太字チェック
            if run.bold:
                style += "font-weight:bold;"

            # 文字色チェック
            if run.font.color and run.font.color.rgb:
                color = run.font.color.rgb
                style += f"color:#{color};"

            # スタイルを適用してHTMLに変換
            paragraph_html += f"<span style='{style}'>{text}</span>"

        # 段落ごとに改行を追加
        html_content += f"<p>{paragraph_html}</p>"

    return html_content


if __name__ == "__main__":
    input_word_file = "./input_h.docx"
    output_word_file = "./output.docx"
    process_word_file(input_word_file, output_word_file)
